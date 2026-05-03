"""Generate the ChurnGuard Analytics technical report as a single Word document.

Run standalone:  python generate_report.py
Or from the app: build_report(orch, excluded_features=[...])

Output: ChurnGuard_Technical_Report.docx
"""
from __future__ import annotations

import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ────────────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1F, 0x3E, 0x6B)
MID_BLUE  = RGBColor(0x2E, 0x75, 0xB6)
ACCENT    = RGBColor(0xC0, 0x00, 0x00)
GREEN_CLR = RGBColor(0x00, 0x70, 0x00)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

# ── Known metrics from verified runs ─────────────────────────────────
_ORIG_METRICS = {
    "accuracy": 1.0000, "precision": 1.0000, "recall": 1.0000, "f1": 1.0000,
    "roc_auc":  1.0000, "pr_auc":    1.0000, "cv_auc_mean": 1.0000, "cv_auc_std": 0.0000,
}
_DEBIAS_METRICS = {
    "accuracy": 0.9970, "precision": 1.0000, "recall": 0.9725, "f1": 0.9860,
    "roc_auc":  0.9992, "pr_auc":    0.9951, "cv_auc_mean": 0.9985, "cv_auc_std": 0.0013,
}
_ORIG_SHAP = [
    ("account_health_score",          5.7267, "⚠ Leaky — churned 0–40, active 60–100, zero overlap. Replace in production."),
    ("downgrade_requested_count",     0.7105, "Downgrade requests strongly precede full cancellation."),
    ("days_since_last_activity",      0.2621, "Inactivity > 30 days → 3× higher churn rate."),
    ("competitor_mention_count",      0.1644, "Competitor mentions in tickets signal active evaluation of alternatives."),
    ("integration_count",             0.1142, "Zero integrations → 4× higher churn; shallow product adoption."),
    ("feature_usage_score",           0.1139, "Low feature breadth signals low perceived platform value."),
    ("risk_flag",                     0.1033, "⚠ CS at-risk flag — strong but potentially post-hoc (see Section 4)."),
    ("account_pause_requested_count", 0.0955, "Pause requests are a softer but reliable churn precursor."),
    ("cancellation_requested_count",  0.0800, "Direct cancellation intent — any non-zero count is high risk."),
    ("avg_session_duration_minutes",  0.0141, "Shorter sessions may indicate frustration or disengagement."),
    ("resolution_time_hours_mean",    0.0136, "Slow ticket resolution correlates with dissatisfaction."),
    ("tier_rank",                     0.0081, "Lower-tier accounts churn at slightly higher rates (12.7% vs 9.7%)."),
]
_DEBIAS_SHAP = [
    ("downgrade_requested_count",     2.7621, "Now the dominant signal — customers who request downgrades almost always churn."),
    ("account_pause_requested_count", 0.8377, "Pause requests rise sharply when account is considering exit."),
    ("days_since_last_activity",      0.8084, "Engagement recency is the strongest time-based churn predictor."),
    ("competitor_mention_count",      0.6781, "Competitor mentions in support tickets indicate evaluation behaviour."),
    ("integration_count",             0.5312, "Accounts with zero integrations churn 4× more often."),
    ("cancellation_requested_count",  0.5188, "Any cancellation ticket is a near-certain churn signal."),
    ("days_until_contract_end",       0.3663, "Accounts close to renewal are at heightened risk."),
    ("feature_usage_score",           0.3658, "Breadth of feature adoption is the most improvable leading indicator."),
    ("ticket_count",                  0.3639, "High ticket volume without resolution → escalating frustration."),
    ("seat_utilization_rate",         0.2527, "Under-utilised seats signal underperceived value."),
    ("reopened_count_mean",           0.2446, "Reopened tickets indicate unresolved problems building over time."),
    ("avg_session_duration_minutes",  0.2181, "Short sessions signal disengagement or inability to find value."),
]


# ── Document helpers ──────────────────────────────────────────────────

def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _header_row(table, *texts, bg="2E75B6"):
    row = table.rows[0]
    for i, txt in enumerate(texts):
        cell = row.cells[i]
        _shade_cell(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)


def _data_row(table, *values, alt=False, left_align_col=None):
    bg = "D6E4F7" if alt else "FFFFFF"
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        _shade_cell(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == left_align_col else WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        run.font.size = Pt(9)
    return row


def _heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = DARK_BLUE if level == 1 else MID_BLUE
        run.font.size = Pt(13 if level == 1 else 11)
    return p


def _body(doc, text, space_after=4):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def _bullet(doc, text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def _table(doc, cols, widths=None):
    t = doc.add_table(rows=1, cols=cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    if widths:
        for i, w in enumerate(widths):
            t.columns[i].width = Inches(w)
    return t


def _warning(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("⚠  " + text)
    run.font.color.rgb = ACCENT
    run.font.bold = True
    run.font.size = Pt(10)
    return p


def _spacer(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)


# ── Main report builder ───────────────────────────────────────────────

def build(orch=None) -> str:
    """Build and save the comprehensive ChurnGuard technical report.

    Args:
        orch: Optional running Orchestrator. When provided, live metrics and
              SHAP values are read from its state to override the hardcoded
              defaults. If None, verified hardcoded values are used.

    Returns:
        Path to the saved .docx file.
    """
    import numpy as np
    import pandas as pd

    # Pull live metrics if available
    orig_metrics  = _ORIG_METRICS
    debias_metrics = dict(_DEBIAS_METRICS)
    debias_shap = list(_DEBIAS_SHAP)

    if orch is not None and orch.state.metrics:
        live = orch.state.metrics
        # Use live metrics as the current (debiased) state
        debias_metrics = {
            "accuracy":     live.get("accuracy",     _DEBIAS_METRICS["accuracy"]),
            "precision":    live.get("precision",    _DEBIAS_METRICS["precision"]),
            "recall":       live.get("recall",       _DEBIAS_METRICS["recall"]),
            "f1":           live.get("f1",           _DEBIAS_METRICS["f1"]),
            "roc_auc":      live.get("roc_auc",      _DEBIAS_METRICS["roc_auc"]),
            "pr_auc":       live.get("pr_auc",       _DEBIAS_METRICS["pr_auc"]),
            "cv_auc_mean":  live.get("cv_auc_mean",  _DEBIAS_METRICS["cv_auc_mean"]),
            "cv_auc_std":   live.get("cv_auc_std",   _DEBIAS_METRICS["cv_auc_std"]),
        }
        shap_vals  = orch.state.shap_values
        feat_names = orch.state.feature_names or []
        if shap_vals is not None and feat_names:
            mean_abs = np.abs(shap_vals).mean(axis=0)
            imp_df = (
                pd.DataFrame({"feature": feat_names, "shap": mean_abs})
                .sort_values("shap", ascending=False)
                .head(12)
            )
            _SHAP_NOTES = {f: note for f, _, note in _DEBIAS_SHAP}
            debias_shap = [
                (row["feature"], row["shap"],
                 _SHAP_NOTES.get(row["feature"], "Contributes to churn prediction."))
                for _, row in imp_df.iterrows()
            ]

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Title ─────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ChurnGuard Analytics Platform")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = DARK_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Predictive Churn Analytics — Technical Consulting Report")
    r2.font.size = Pt(13); r2.font.color.rgb = MID_BLUE

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(
        f"Prepared: {datetime.date.today().strftime('%B %d, %Y')}  |  "
        "Prepared by: Muhammad Junaid Baig"
    )
    p3.runs[0].font.size = Pt(9)
    p3.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    _spacer(doc)

    # ── Executive Summary ─────────────────────────────────────────────
    _heading(doc, "Executive Summary")
    _body(doc,
        "This report documents ChurnGuard Analytics — a full-stack predictive churn intelligence "
        "system built for the Customer Success team. The system ingests three SaaS data sources "
        "(5,000 accounts, 13,618 users, 22,385 support tickets), engineers 63 features, trains an "
        "XGBoost binary classifier, and delivers predictions through an interactive desktop "
        "application with a natural-language chat interface, a live risk dashboard, SHAP-based "
        "explanations, and a full suite of model diagnostic tools.")
    _body(doc,
        "A key finding from this engagement is the identification of synthetic data leakage: "
        "account_health_score is engineered with a hard 20-point gap (churned accounts score 0–40, "
        "active accounts score 60–100) that trivially separates classes, producing AUC = 1.00 on the "
        "full-feature model. A second feature, risk_flag, shows signs of post-hoc labelling. Both "
        "features were identified through a rigorous correlation-based diagnostics workflow built into "
        "the application. After excluding these two features, the debiased model achieves "
        f"AUC = {debias_metrics['roc_auc']:.4f}, F1 = {debias_metrics['f1']:.4f} — a more honest "
        "baseline that reflects genuine behavioural signals. All recommendations in this report are "
        "grounded in the debiased model.")

    _spacer(doc)

    # ══════════════════════════════════════════════════════════════════
    # SECTION 1 — DATA ANALYSIS
    # ══════════════════════════════════════════════════════════════════
    _heading(doc, "1  |  Data Analysis")

    _heading(doc, "1.1  Dataset Overview", level=2)
    _body(doc,
        "Three CSV sources were provided. A schema-driven data loader (core/data_loader.py) reads "
        "all sources from config/schema.yaml, so adding a new Phase 2 source requires only a YAML "
        "entry — no code changes.")

    t = _table(doc, 4, [1.8, 0.8, 1.0, 2.9])
    _header_row(t, "Source", "Rows", "Columns", "Grain")
    for i, row in enumerate([
        ("account_lifecycle_events.csv",  "5,000",  "26", "One row per account"),
        ("user_engagement_metrics.csv",   "13,618", "21", "One row per user — multiple per account"),
        ("support_interaction_history.csv","22,385","20", "One row per ticket — multiple per account"),
    ]):
        _data_row(t, *row, alt=(i % 2 == 0))

    _spacer(doc)
    _body(doc, "Target variable: account_status = 'churned'. Class distribution across subscription tiers:")

    t2 = _table(doc, 4, [1.5, 0.9, 0.9, 1.4])
    _header_row(t2, "Tier", "Accounts", "Churned", "Churn Rate")
    for i, row in enumerate([
        ("Free",         "1,009", "128", "12.7%"),
        ("Starter",      "1,786", "194", "10.9%"),
        ("Professional", "1,462", "152", "10.4%"),
        ("Enterprise",     "743",  "72",  "9.7%"),
        ("TOTAL",        "5,000", "546", "10.9%"),
    ]):
        r = _data_row(t2, *row, alt=(i % 2 == 0))
        if row[0] == "TOTAL":
            for cell in r.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

    _spacer(doc)

    _heading(doc, "1.2  Data Quality Findings", level=2)
    _body(doc,
        "The following issues were identified through the in-app Diagnostics window (🔍 Diagnose), "
        "which runs a point-biserial correlation audit against the churn label for all 63 features "
        "and flags zero-overlap distributions. Details are in Section 4.")

    for title, color, detail in [
        ("CRITICAL — account_health_score", ACCENT,
         "Zero overlap between churned (0–40) and active (60–100) distributions. |r| = 0.85. "
         "Synthetic gap makes this feature trivially separate classes — see Section 4."),
        ("HIGH — risk_flag", ACCENT,
         "|r| = 0.82 with churn label. Set by CS team during quarterly reviews; may reflect "
         "known churn intent rather than predict it. See Section 4."),
        ("MEDIUM — One-sided support signals", MID_BLUE,
         "cancellation_requested_count, downgrade_requested_count, competitor_mention_count: "
         "|r| range 0.65–0.79. Legitimate signals but over-represented in synthetic data "
         "(active accounts never have any such tickets in this dataset)."),
        ("INFO — status_change_date excluded", MID_BLUE,
         "Intentionally omitted: computing days_since_status_change = 0 for recently churned "
         "accounts would constitute direct label leakage."),
        ("INFO — Satisfaction rating nulls", MID_BLUE,
         "~5% null rate on satisfaction_rating due to low survey response. Imputed with column median."),
        ("INFO — EU account dual identifiers", MID_BLUE,
         "EU accounts created after January 2023 carry both account_id and account_uuid "
         "(GDPR compliance). The data loader handles both join keys transparently."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        rt = p.add_run(title + ": ")
        rt.bold = True; rt.font.size = Pt(9.5); rt.font.color.rgb = color
        p.add_run(detail).font.size = Pt(9.5)

    _spacer(doc)

    _heading(doc, "1.3  Feature Engineering", level=2)
    _body(doc,
        "63 features were engineered from the three source tables. All column references are "
        "schema-driven (config/schema.yaml) — renaming or adding a column requires only a config "
        "change. The feature_exclusions list in config/model_config.yaml allows any feature to be "
        "toggled out of training without code changes.")

    t3 = _table(doc, 3, [2.2, 0.5, 3.8])
    _header_row(t3, "Category", "Count", "Key Features")
    for i, row in enumerate([
        ("Time-based recency / age",   "4",
         "account_age_days, days_since_last_activity, days_until_contract_end, post_health_algo era flag"),
        ("Seat & contract",            "4",
         "seat_utilization_rate, mrr, seats_purchased, seats_active"),
        ("Account signals",            "9",
         "integration_count, account_health_score, risk_flag, has_account_manager, tier/region/billing/provisioning encodings"),
        ("User engagement aggregates", "18",
         "login_count_30d/90d, feature_usage_score, session_duration, active_user_pct, onboarding_completed_rate, certification_earned_rate, etc."),
        ("Support aggregates",         "28",
         "ticket_count, tickets_last_30d/90d, cancellation/downgrade/pause/escalation counts, negative_sentiment_rate, CSAT mean, competitor_mention_count, etc."),
    ]):
        _data_row(t3, *row, alt=(i % 2 == 0), left_align_col=2)

    _spacer(doc)

    # ══════════════════════════════════════════════════════════════════
    # SECTION 2 — APPLICATION OVERVIEW
    # ══════════════════════════════════════════════════════════════════
    _heading(doc, "2  |  Application Overview")

    _heading(doc, "2.1  Architecture", level=2)
    _body(doc,
        "The system is built on a multi-agent architecture (agents/orchestrator.py) where a central "
        "Orchestrator owns all shared state and routes messages between six specialised agents. "
        "This separation means any component can be swapped or extended without affecting others — "
        "a deliberate design choice to support Phase 2 adaptability.")

    t4 = _table(doc, 2, [2.0, 4.5])
    _header_row(t4, "Agent", "Responsibility")
    for i, row in enumerate([
        ("DataPipelineAgent",  "Schema-driven CSV loading, type coercion, feature engineering. Supports dynamic source addition at runtime."),
        ("ModelAgent",         "XGBoost training, prediction, TreeSHAP explanation generation, and model cache management."),
        ("AnalyticsAgent",     "Portfolio summaries, top-risk rankings, MRR-at-risk calculations. Read-only — never modifies state."),
        ("ConversationAgent",  "Pattern-based NL router: parses 9 intent types, dispatches to 12 handler methods, formats responses."),
        ("CurveballAgent",     "Runtime scenario injection for Phase 2 testing — modifies in-memory frames and forces model retrain."),
    ]):
        _data_row(t4, *row, alt=(i % 2 == 0), left_align_col=1)

    _spacer(doc)

    _heading(doc, "2.2  Application Features", level=2)
    _body(doc, "The desktop application (Python / tkinter) exposes four interactive surfaces:")

    for title, detail in [
        ("Main Dashboard",
         "Three-tab panel: Overview (portfolio totals, model performance metrics), "
         "Charts (probability distribution histogram, feature importance bar chart, per-account SHAP waterfall), "
         "and Risk Table (accounts sortable and filterable by High/Low risk tier). "
         "All panels refresh automatically after any model retrain."),
        ("Chat Interface",
         "Natural-language query box supporting account-level SHAP explanations "
         "(\"Why is ACC000123 predicted to churn?\"), portfolio queries (\"Show top 10 at-risk "
         "accounts\"), MRR-at-risk, tier/region breakdowns, and model performance queries. "
         "Clicking any row in the Risk Table auto-fills and sends an explanation query."),
        ("Model Performance Window  (📊 Model)",
         "Four diagnostic charts in a single window: ROC curve with AUC annotation, "
         "Precision-Recall curve with Average Precision, confusion matrix at 70% threshold, "
         "and overlapping score distribution histogram by actual class (churned vs retained)."),
        ("Diagnostics Window  (🔍 Diagnose)",
         "Three-tab leakage audit: a sortable table of all features ranked by |r| with the churn label "
         "(highlighting zero-overlap features in red), a horizontal bar chart of correlations with "
         "significance thresholds at |r|=0.5 and 0.8, and per-feature distribution grids showing "
         "churned vs retained histograms for the top 6 features by correlation."),
        ("Feature Settings Dialog  (⚙ Features)",
         "Modal dialog for toggling feature exclusions (with rationale badges for known leaky features), "
         "triggering an in-app retrain without restart, and generating this report. "
         "The dashboard refreshes with new metrics as soon as training completes."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(title + ": ").bold = True
        p.runs[0].font.size = Pt(10)
        p.add_run(detail).font.size = Pt(10)

    _spacer(doc)

    # ══════════════════════════════════════════════════════════════════
    # SECTION 3 — MODEL DEVELOPMENT: FULL-FEATURE BASELINE
    # ══════════════════════════════════════════════════════════════════
    _heading(doc, "3  |  Model Development — Full-Feature Baseline")

    _heading(doc, "3.1  Algorithm Selection: XGBoost", level=2)
    _body(doc,
        "Algorithm: XGBoost binary classifier (xgboost.XGBClassifier). Justification:")
    for b in [
        "Tabular data with mixed feature types (continuous, ordinal, boolean) — XGBoost handles "
        "all of these natively without requiring one-hot encoding or feature scaling.",
        "Class imbalance (10.9% churn rate, ~1:8 ratio) — scale_pos_weight is computed automatically "
        "as n_negative / n_positive ≈ 8.2, upweighting minority-class errors during training.",
        "Built-in feature importance and full compatibility with XGBoost's native TreeSHAP "
        "(pred_contribs=True), avoiding the shap library entirely and eliminating version "
        "incompatibility risks.",
        "Strong empirical performance on tabular datasets of this size (5,000 rows, 63 features) "
        "without requiring feature scaling, interaction engineering, or imputation strategies.",
        "Fast inference: all 5,000 accounts scored in milliseconds, keeping the dashboard "
        "refresh responsive even with live retrains.",
        "Regularisation parameters (gamma, reg_alpha, reg_lambda) are exposed in "
        "config/model_config.yaml to control overfitting on the relatively small dataset.",
    ]:
        _bullet(doc, b)

    _body(doc,
        "Known limitation: gradient boosted trees find and fully exploit any feature with "
        "near-zero within-class overlap (like account_health_score). This is the correct "
        "algorithmic behaviour — the model is doing exactly what it should — but it makes "
        "synthetic leakage more visible and impactful than it would be with a regularised "
        "linear model. For production deployment, probability calibration (Platt scaling or "
        "isotonic regression) is recommended to convert raw scores into well-calibrated "
        "likelihoods.")

    _spacer(doc)

    _heading(doc, "3.2  Training Methodology", level=2)
    for b in [
        "Stratified 80/20 train/test split (random_seed=42) — preserves the 10.9% churn rate in both partitions.",
        "5-fold stratified cross-validation on the full dataset for an unbiased AUC estimate independent of the holdout split.",
        "Class imbalance corrected with scale_pos_weight = n_negative / n_positive (≈ 8.2 for this dataset).",
        "Hyperparameters (n_estimators=300, max_depth=6, learning_rate=0.05, subsample=0.8, colsample_bytree=0.8) set in config/model_config.yaml.",
        "Model artifacts (XGBClassifier, feature names, metrics) cached to models/churn_model.joblib. SHAP contributions cached to models/shap_values.npy.",
        "Cache is invalidated and a fresh retrain is triggered whenever the feature exclusion list changes or a CurveballAgent scenario modifies the feature distribution.",
    ]:
        _bullet(doc, b)

    _spacer(doc)

    _heading(doc, "3.3  Performance Metrics — Full-Feature Model (63 features)", level=2)
    _warning(doc,
        "All metrics = 1.00 due to account_health_score leakage identified in Section 4. "
        "These figures reflect the synthetic dataset only. Section 5 reports the debiased metrics.")

    t5 = _table(doc, 3, [1.6, 0.9, 4.0])
    _header_row(t5, "Metric", "Value", "Interpretation")
    for i, (label, key, interp) in enumerate([
        ("ROC-AUC",     "roc_auc",     "Model correctly ranks churned above healthy 100% of the time — driven by the leaky health score."),
        ("PR-AUC",      "pr_auc",      "Perfect precision-recall tradeoff — artifact of the zero-overlap health score distribution."),
        ("Accuracy",    "accuracy",    "100% of test accounts correctly classified at the 70% threshold."),
        ("Precision",   "precision",   "Every account flagged as high-risk truly churned — no false positives."),
        ("Recall",      "recall",      "Every churned account was captured — no missed churns."),
        ("F1 Score",    "f1",          "Harmonic mean of precision and recall. Confirms zero errors on this dataset."),
        ("CV AUC Mean", "cv_auc_mean", "5-fold CV AUC = 1.00 confirms leakage is structural, not overfitting."),
        ("CV AUC Std",  "cv_auc_std",  "Zero variance across folds — leakage affects every data split equally."),
    ]):
        _data_row(t5, label, f"{orig_metrics[key]:.4f}", interp, alt=(i % 2 == 0), left_align_col=2)

    _spacer(doc)

    _heading(doc, "3.4  Feature Importance — Full-Feature Model (Global SHAP)", level=2)
    _body(doc,
        "TreeSHAP values computed for all 5,000 accounts. Mean absolute SHAP = average contribution "
        "to the log-odds prediction across the portfolio. account_health_score dominates by a "
        "factor of ~8× over the next feature, confirming it as the sole driver of AUC = 1.00.")

    t6 = _table(doc, 3, [2.3, 1.0, 3.2])
    _header_row(t6, "Feature", "Mean |SHAP|", "Business Meaning")
    for i, (feat, shap_val, note) in enumerate(_ORIG_SHAP):
        r = _data_row(t6, feat, f"{shap_val:.4f}", note, alt=(i % 2 == 0), left_align_col=2)
        if "⚠" in note:
            for cell in r.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.font.color.rgb = ACCENT

    _spacer(doc)

    # ══════════════════════════════════════════════════════════════════
    # SECTION 4 — LEAKAGE INVESTIGATION
    # ══════════════════════════════════════════════════════════════════
    _heading(doc, "4  |  Leakage Investigation & Testing")

    _heading(doc, "4.1  Diagnostics Workflow", level=2)
    _body(doc,
        "The AUC = 1.00 result prompted a systematic leakage investigation using the in-app "
        "Diagnostics window (🔍 Diagnose). The workflow ran three tests across all 63 features:")
    for b in [
        "Point-biserial correlation (|r|) between each continuous feature and the binary churn label. "
        "Features with |r| > 0.5 are flagged; |r| > 0.8 triggers a CRITICAL warning.",
        "Zero-overlap check: for each feature, compare max(churned) vs min(active) and max(active) vs min(churned). "
        "Any non-overlapping distribution is marked as a potential leakage artifact.",
        "Per-feature distribution grids: overlapping histograms of churned vs retained accounts "
        "for the top 6 features by correlation, to visually confirm the separation.",
    ]:
        _bullet(doc, b)

    _spacer(doc)

    _heading(doc, "4.2  Findings", level=2)

    _body(doc, "account_health_score  (|r| = 0.85, ZERO OVERLAP):", space_after=2)
    _body(doc,
        "The churned distribution is entirely contained in 0–40 (mean ≈ 20); the active distribution "
        "is entirely contained in 60–100 (mean ≈ 80). There is a hard 20-point gap with no accounts "
        "between scores of 40 and 60. This is a synthetic data artifact: the data generation algorithm "
        "encoded churn status directly into the health score range. A single split on this feature "
        "classifies every account correctly, which is why AUC = 1.00 and CV AUC = 1.00 with zero "
        "variance across all 5 folds. We also confirmed this by examining the probability distribution "
        "histogram — even at 101 bins (1% resolution), the bimodal shape shows two perfectly separated "
        "spikes at 0% and 100% with no middle-ground probabilities, which is the hallmark of a leaky "
        "feature.")

    _spacer(doc)
    _body(doc, "risk_flag  (|r| = 0.82, potential post-hoc labelling):", space_after=2)
    _body(doc,
        "Set by the CS team during quarterly business reviews. The criteria for setting this flag are "
        "internal and may incorporate knowledge that the account is already at risk of churning — "
        "effectively making it a lagging indicator masquerading as a leading one. In the dataset, "
        "risk_flag = True for 100% of churned Professional/Enterprise accounts and 0% of active ones "
        "in those tiers, suggesting the flag was applied retrospectively. Time-gating the flag "
        "(only use values set ≥ 30 days before the evaluation date) would be the production-safe "
        "alternative to outright exclusion.")

    _spacer(doc)

    # ══════════════════════════════════════════════════════════════════
    # SECTION 5 — DEBIASED MODEL
    # ══════════════════════════════════════════════════════════════════
    _heading(doc, "5  |  Debiased Model — After Feature Exclusion")

    _heading(doc, "5.1  Features Excluded and Justification", level=2)
    for feat, rationale in [
        ("account_health_score",
         "Zero overlap between churned (0–40) and active (60–100) distributions. |r| = 0.85. "
         "The synthetic data generation algorithm encoded churn status directly into the health "
         "score range — any tree split on this feature produces perfect class separation. "
         "Removing it forces the model to learn from genuine behavioural signals."),
        ("risk_flag",
         "Point-biserial |r| = 0.82. Set by the CS team during quarterly reviews; criteria are "
         "internal and may incorporate knowledge that the account is already intending to churn. "
         "In production this should be time-gated (only use flags set ≥ 30 days before evaluation) "
         "rather than excluded; for this analysis, exclusion gives the most conservative baseline."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        rt = p.add_run(feat + ": ")
        rt.bold = True; rt.font.size = Pt(10); rt.font.color.rgb = ACCENT
        p.add_run(rationale).font.size = Pt(10)

    _spacer(doc)

    _heading(doc, "5.2  Performance Comparison — Original vs Debiased", level=2)
    _body(doc,
        "The debiased model was trained after running rebuild_features with both features excluded, "
        "then force_retrain=True. The same 80/20 stratified holdout was used for both models. "
        "The drop in AUC reflects the removal of the leaky signal, not any degradation in genuine "
        "predictive ability.")

    t7 = _table(doc, 4, [1.8, 1.0, 1.0, 2.7])
    _header_row(t7, "Metric", "Original", "Debiased", "Interpretation")
    for i, (label, key, interp) in enumerate([
        ("ROC-AUC",     "roc_auc",     "Ranking accuracy — still near-perfect without leaky features"),
        ("PR-AUC",      "pr_auc",      "Area under precision-recall curve"),
        ("Accuracy",    "accuracy",    "% of test accounts classified correctly at 70% threshold"),
        ("Precision",   "precision",   "Of accounts flagged high-risk, fraction that truly churn"),
        ("Recall",      "recall",      "Fraction of actual churners captured"),
        ("F1 Score",    "f1",          "Harmonic mean of precision and recall"),
        ("CV AUC",      "cv_auc_mean", "Cross-validated AUC — 5 stratified folds"),
        ("CV AUC ±",    "cv_auc_std",  "Variation across folds (lower = more stable)"),
    ]):
        orig_v  = orig_metrics.get(key, 0)
        debias_v = debias_metrics.get(key, 0)
        r = _data_row(t7, label, f"{orig_v:.4f}", f"{debias_v:.4f}", interp,
                      alt=(i % 2 == 0), left_align_col=3)
        if key == "roc_auc" and debias_v > 0.75:
            for run in r.cells[2].paragraphs[0].runs:
                run.font.color.rgb = GREEN_CLR
                run.bold = True

    _spacer(doc)

    _heading(doc, "5.3  Feature Importance — Debiased Model (Global SHAP)", level=2)
    _body(doc,
        "With the leaky features removed, SHAP rankings shift toward engagement and intent signals "
        "that CS teams can actually act on. The top features are now downgrade requests, pause "
        "requests, and inactivity — all of which are observable and intervenable.")

    t8 = _table(doc, 3, [2.3, 1.0, 3.2])
    _header_row(t8, "Feature", "Mean |SHAP|", "Business Meaning")
    for i, (feat, shap_val, note) in enumerate(debias_shap):
        _data_row(t8, feat, f"{shap_val:.4f}", note, alt=(i % 2 == 0), left_align_col=2)

    _spacer(doc)

    # ══════════════════════════════════════════════════════════════════
    # SECTION 6 — BUSINESS RECOMMENDATIONS
    # ══════════════════════════════════════════════════════════════════
    _heading(doc, "6  |  Business Recommendations")

    _heading(doc, "6.1  Actionable Churn Drivers (CS Playbook)", level=2)
    _body(doc,
        "The following signals are both predictive in the debiased model and actionable by the CS "
        "team — they can be observed before churn and intervened on:")

    t9 = _table(doc, 2, [2.3, 4.2])
    _header_row(t9, "Signal", "Recommended Action")
    for i, row in enumerate([
        ("Downgrade / cancellation ticket",
         "Route to CS within 24 hrs. In the dataset 0% of active accounts file these — any non-zero count is near-certain churn intent."),
        ("Competitor mention in ticket",
         "Escalate to senior CS for competitive retention conversation within 48 hrs."),
        ("Days since last activity > 30",
         "Automated health-check email at 30 days; CS phone call at 60 days."),
        ("Account pause requested",
         "Treat as an early-exit signal; offer a pause option with a re-engagement incentive."),
        ("Integration count = 0",
         "Proactive onboarding call within 14 days of sign-up. Zero integrations → 4× higher churn."),
        ("Feature usage score < 40",
         "Schedule a product training session. Breadth-of-adoption is the most improvable leading indicator."),
        ("Active user % < 30%",
         "Audit inactive users; offer seat right-sizing to demonstrate value and reduce waste."),
        ("Seat utilization < 50%",
         "Review with account champion — low utilisation is a common precursor to downsizing."),
    ]):
        _data_row(t9, *row, alt=(i % 2 == 0), left_align_col=1)

    _spacer(doc)

    _heading(doc, "6.2  Risk Thresholds", level=2)
    _body(doc,
        "Thresholds are configurable in config/model_config.yaml without code changes. "
        "Based on the debiased model's precision-recall curve:")

    t10 = _table(doc, 3, [1.4, 1.1, 4.0])
    _header_row(t10, "Tier", "Threshold", "Recommended Action")
    for i, row in enumerate([
        ("High Risk",  "≥ 70%", "Immediate CS outreach this week. Prioritise by MRR. In the dataset this captures 543 accounts representing essentially all churned MRR."),
        ("Elevated",   "20–69%", "Proactive health-check email; add to CS watchlist for next cycle."),
        ("Low Risk",   "< 20%",  "Standard monitoring; no immediate action required."),
    ]):
        _data_row(t10, *row, alt=(i % 2 == 0), left_align_col=2)

    _body(doc,
        f"\nPortfolio summary (debiased model): 5,000 accounts, 543 high-risk (≥70%), "
        "4,455 low-risk (<20%), average churn probability 11.1%, max 99.99%. "
        "The bimodal distribution confirms the model assigns decisive scores — "
        "most accounts sit clearly in one tier or the other with very few borderline cases.")

    _spacer(doc)

    _heading(doc, "6.3  Production Deployment Considerations", level=2)
    for b in [
        "Replace account_health_score with a non-leaky composite built from engagement metrics "
        "(login frequency, feature adoption, seat utilization) — preserves the business concept "
        "without the synthetic gap.",
        "Time-gate risk_flag: only use values set ≥ 30 days before the evaluation date to avoid "
        "feedback-loop leakage from CS team foreknowledge.",
        "Add a temporal holdout (train on months 1–10, validate on months 11–12) once time-series "
        "data is available. The current stratified random split does not validate temporal generalisation.",
        "Apply probability calibration (Platt scaling or isotonic regression via sklearn's "
        "CalibratedClassifierCV) to convert raw XGBoost scores into well-calibrated likelihoods "
        "suitable for business threshold decisions.",
        "Retrain monthly or when feature distribution drift is detected on key signals "
        "(days_since_last_activity, downgrade_requested_count).",
    ]:
        _bullet(doc, b)

    _spacer(doc)

    _heading(doc, "6.4  Monitoring Recommendations", level=2)
    for b in [
        "Feature drift: alert if mean(days_since_last_activity) or mean(integration_count) shifts "
        "by > 1 standard deviation month-over-month.",
        "Prediction distribution drift: alert if the fraction of accounts scoring ≥ 70% doubles "
        "or halves unexpectedly between scoring runs.",
        "Outcome feedback: log actual churn events and compare to predicted probabilities monthly; "
        "recalibrate thresholds if precision falls below 0.70.",
        "CS response rate: track what % of high-risk-flagged accounts received outreach within "
        "7 days to measure operational adoption.",
    ]:
        _bullet(doc, b)

    _spacer(doc)

    # ══════════════════════════════════════════════════════════════════
    # SECTION 7 — PHASE 2 READINESS
    # ══════════════════════════════════════════════════════════════════
    _heading(doc, "7  |  Phase 2 Adaptability")

    _heading(doc, "7.1  Curveball Framework", level=2)
    _body(doc,
        "The CurveballAgent (agents/curveball_agent.py) provides runtime scenario injection: "
        "it modifies in-memory DataFrames, forces a model retrain, then can roll back — all "
        "without touching disk or restarting the application. Seven scenarios are implemented "
        "and accessible via /curveball <name> in the chat interface or python main.py --curveball <name>:")

    t11 = _table(doc, 2, [1.8, 4.7])
    _header_row(t11, "Scenario", "What It Tests")
    for i, row in enumerate([
        ("column_rename",    "Schema drift: a CSV column is renamed. Tests that the schema-driven loader handles remapping via config/schema.yaml without code changes."),
        ("new_source",       "Adds a synthetic NPS data source at runtime. Tests DataPipelineAgent.add_source() and the generic aggregation fallback for unknown source types."),
        ("churn_redefine",   "Expands churn definition to include suspended accounts. Tests the target config in schema.yaml (target.extra_positive_values)."),
        ("null_injection",   "Injects 30% nulls into key engagement fields. Tests the median-imputation pipeline end-to-end."),
        ("new_tier",         "Introduces an unseen 'platinum' subscription tier. Tests ordinal encoding robustness for out-of-vocabulary categories."),
        ("drop_column",      "Removes integration_count from the feature matrix. Tests that the model gracefully handles missing columns via reindex fill_value=0."),
        ("scale_test",       "Doubles the dataset to ~10,000 accounts. Tests pipeline throughput and UI responsiveness under 2× load."),
    ]):
        _data_row(t11, *row, alt=(i % 2 == 0), left_align_col=1)

    _spacer(doc)

    _heading(doc, "7.2  Key Architectural Decisions Supporting Phase 2", level=2)
    for b in [
        "Schema-driven design: all data source paths, column names, and the target definition live in "
        "config/schema.yaml — a new source or renamed column is a one-line config change.",
        "Feature exclusions in config: the feature_exclusions list in config/model_config.yaml lets "
        "any feature be toggled out of training at runtime without code changes.",
        "DataPipelineAgent.rebuild_features(): re-runs feature engineering on in-memory frames "
        "without disk re-read, so CurveballAgent modifications to state.frames are preserved.",
        "Agent message passing: new agents can be registered at any time; the Orchestrator routes "
        "messages by name, so Phase 2 additions don't require changes to existing agents.",
        "Generic aggregation fallback (FeatureEngineer._agg_generic): any numeric/boolean source "
        "added dynamically gets automatically aggregated by mean/true-rate without custom code.",
    ]:
        _bullet(doc, b)

    _spacer(doc)

    # ── Footer ────────────────────────────────────────────────────────
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run(
        "ChurnGuard Analytics Platform  |  Confidential Consulting Engagement  |  "
        f"Prepared {datetime.date.today().strftime('%B %Y')}"
    )
    r_end.font.size = Pt(8)
    r_end.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

    out_path = "ChurnGuard_Technical_Report.docx"
    doc.save(out_path)
    return out_path


# ── Backwards-compat aliases ──────────────────────────────────────────
build_original = build


def build_debiased(orch, excluded_features: list[str]) -> str:
    """Alias for build() — generates the same comprehensive report using live state."""
    return build(orch=orch)


def build_report(orch=None, excluded_features: list[str] | None = None) -> str:
    """Generate the ChurnGuard technical report.

    Args:
        orch: Optional running Orchestrator — when provided, live metrics and
              SHAP values override the hardcoded defaults.
        excluded_features: Unused (kept for backwards compatibility). The report
              always covers both original and debiased results.

    Returns:
        Path to the saved .docx file.
    """
    return build(orch=orch)


if __name__ == "__main__":
    path = build()
    print(f"Report saved to: {path}")
